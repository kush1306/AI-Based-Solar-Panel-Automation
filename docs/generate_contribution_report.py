"""Generate Member 3 Project Contribution Report as .docx"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION

OUTPUT = Path(__file__).resolve().parent / "Member3_Work_Document.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
H3_SIZE = Pt(13)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(doc, text, *, bold=False, italic=False, size=BODY_SIZE, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, H2_SIZE), bold=True)
    if level == 1:
        p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_figure_placeholder(doc, title, height_inches=3.0):
    add_paragraph(doc, title, bold=True, size=H3_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.0)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n[ Insert screenshot here ]\n\n")
    set_run_font(run, italic=True)
    p.paragraph_format.space_before = Pt(int(height_inches * 30))
    p.paragraph_format.space_after = Pt(int(height_inches * 30))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    cap_run = cap.add_run(f"Figure: {title}")
    set_run_font(cap_run, size=Pt(11), italic=True)
    doc.add_paragraph()


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=Pt(10))
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    add_paragraph(
        doc,
        "Note: Open this document in Microsoft Word and press Ctrl+A, then F9 to update the Table of Contents.",
        italic=True,
        size=Pt(10),
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )


def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        add_page_number_footer(section)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # --- Cover Page ---
    for _ in range(6):
        doc.add_paragraph()
    add_paragraph(doc, "WORK DOCUMENT", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc, "Project Contribution Report", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    add_paragraph(
        doc,
        "Smart Solar Energy Optimization System –\nAI-Driven Solar Generation, Demand & Cost Optimization",
        bold=True,
        size=Pt(14),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(doc, "Submitted By", bold=True, size=H2_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "[Your Full Name]", size=H2_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "[Roll Number / Enrollment Number]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_paragraph(doc, "Role: Member 3 – Full Stack Developer", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "Department: [Your Department]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "Institution: [Your Institution Name]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_paragraph(doc, "Under the Guidance of", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "[Faculty Guide Name]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "[Designation]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_paragraph(doc, "Submission Date: [Month Year]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "GitHub Branch: member3-fullstack", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    doc.add_page_break()

    # --- Table of Contents ---
    add_heading(doc, "Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # --- 1. Project Overview ---
    add_heading(doc, "1. Project Overview", level=1)
    add_paragraph(
        doc,
        "The Smart Solar Energy Optimization System is an AI-driven platform designed to monitor solar "
        "energy generation, predict energy demand, optimize solar panel orientation, and support "
        "cost-efficient energy management. The system integrates sensor telemetry, weather intelligence, "
        "battery monitoring, alerting, and analytics into a unified monitoring environment.",
    )
    add_paragraph(
        doc,
        "The project is developed collaboratively by a multidisciplinary team. Member 1 and Member 2 "
        "focus on AI model development for solar position and energy forecasting. Member 4 handles "
        "DevOps and cloud deployment. Member 5 supports data engineering activities. As Member 3, my "
        "individual contribution is limited to full-stack development: designing and connecting the "
        "MySQL database layer, building the FastAPI backend with complete REST APIs, and developing "
        "the Next.js dashboard frontend.",
    )
    add_paragraph(
        doc,
        "My work enables the platform to store operational data reliably, expose structured APIs for "
        "frontend and future AI integration, and present a professional user interface for monitoring "
        "solar plant performance in real time.",
    )

    # --- 2. My Role ---
    add_heading(doc, "2. My Role and Responsibilities", level=1)
    add_paragraph(doc, "As Member 3 – Full Stack Developer, I was responsible for the following deliverables:", bold=False)
    responsibilities = [
        "Database design alignment, table mapping, and MySQL connectivity for the existing schema.",
        "Development of a production-ready FastAPI backend with SQLAlchemy ORM and Pydantic validation.",
        "Implementation of complete CRUD REST APIs for all nine project database tables.",
        "Creation of health check endpoints, mock AI placeholder endpoints, and dashboard aggregation APIs.",
        "Design and development of the Next.js dashboard with responsive enterprise-grade UI.",
        "Configuration of CORS, environment-based settings, logging, and centralized error handling.",
        "Documentation, local run scripts, and DevOps handoff notes for deployment by Member 4.",
        "Version control and push of all deliverables to the member3-fullstack GitHub branch.",
    ]
    for item in responsibilities:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "I did not implement real AI model inference or Docker/AWS deployment, as those responsibilities "
        "belong to other team members. My backend includes mock endpoints so the frontend can be tested "
        "before AI models are integrated.",
    )

    # --- 3. Database ---
    add_heading(doc, "3. Database Design and Implementation", level=1)
    add_heading(doc, "3.1 Database Selection", level=2)
    add_paragraph(
        doc,
        "MySQL 8.x was selected as the relational database management system because it is widely used "
        "in industrial and academic projects, supports structured relational data, provides strong "
        "referential integrity through foreign keys, and integrates efficiently with Python through "
        "SQLAlchemy and PyMySQL. MySQL Workbench was used for visual schema design, EER diagram "
        "creation, and query validation.",
    )
    add_heading(doc, "3.2 Database Design Process", level=2)
    add_paragraph(
        doc,
        "The database design process began with identifying the core entities required for solar "
        "monitoring: weather conditions, solar panels, predictions, energy consumption, batteries, "
        "telemetry readings, alerts, and system logs. I analyzed functional requirements from the "
        "project scope and translated them into normalized relational tables with appropriate primary "
        "keys, foreign keys, and data types.",
    )
    add_paragraph(
        doc,
        "An Entity-Relationship (EER) diagram was prepared in MySQL Workbench to visualize table "
        "relationships before implementation. The final database name is solar_panel_automation. The "
        "backend was developed around the existing schema without modifying table structures, ensuring "
        "compatibility with team data and AI modules.",
    )
    add_figure_placeholder(doc, "Database EER Diagram", height_inches=3.5)
    add_figure_placeholder(doc, "MySQL Workbench Screenshot", height_inches=3.0)

    add_heading(doc, "3.3 Tables and Relationships", level=2)
    add_paragraph(doc, "The database consists of nine tables:", bold=False)
    tables = [
        "weather_data – stores environmental readings (temperature, humidity, GHI, DNI, AQI, city).",
        "solar_panel – stores panel metadata (name, capacity, efficiency, tilt, status).",
        "solar_predictions – stores predicted tilt and expected power linked to panel and weather.",
        "energy_consumption – stores load consumption records with temporal features.",
        "battery – stores battery device information and health percentage.",
        "battery_status – stores SOC, voltage, current, and charging status snapshots.",
        "telemetry – stores live sensor readings from panels and batteries.",
        "alerts – stores system alerts with severity, type, and resolution status.",
        "system_logs – stores application and system event logs.",
    ]
    for t in tables:
        add_bullet(doc, t)

    add_heading(doc, "3.4 Primary Keys and Foreign Keys", level=2)
    add_paragraph(
        doc,
        "Each table uses an auto-increment integer primary key (e.g., weather_id, panel_id, "
        "prediction_id). Foreign key relationships enforce data integrity: solar_predictions references "
        "solar_panel and weather_data; battery_status references battery; telemetry references both "
        "solar_panel and battery; alerts optionally reference solar_panel and battery.",
    )
    add_heading(doc, "3.5 Sample Data and Testing", level=2)
    add_paragraph(
        doc,
        "Sample records were inserted into all tables using MySQL Workbench and validated through "
        "backend API calls. List, create, update, and delete operations were tested for each entity "
        "using Swagger UI. The /health/ready endpoint confirms live database connectivity during runtime.",
    )

    # --- 4. Backend ---
    add_heading(doc, "4. Backend Development", level=1)
    add_heading(doc, "4.1 FastAPI Project Setup", level=2)
    add_paragraph(
        doc,
        "The backend was built using FastAPI with Uvicorn as the ASGI server. The application entry "
        "point is app/main.py, running on port 8000. Configuration is loaded from backend/.env using "
        "Pydantic Settings. Dependencies are listed in requirements.txt and include FastAPI, SQLAlchemy, "
        "PyMySQL, Pydantic, python-dotenv, and Uvicorn.",
    )
    add_heading(doc, "4.2 Project Folder Structure", level=2)
    add_paragraph(
        doc,
        "The backend follows a modular layered architecture separating routes, business logic, data "
        "access, schemas, and configuration:",
    )
    structure_lines = [
        "app/api/ – REST route handlers for each resource",
        "app/crud/ – reusable CRUD operations per database table",
        "app/dependencies/ – database session injection and router factory",
        "app/models/ – SQLAlchemy ORM models mapped to MySQL tables",
        "app/schemas/ – Pydantic Create, Update, and Response schemas",
        "app/services/ – dashboard aggregation and mock AI services",
        "app/core/ – configuration, database engine, logging, and exceptions",
        "app/utils/ – pagination and shared helpers",
    ]
    for line in structure_lines:
        add_bullet(doc, line)
    add_figure_placeholder(doc, "Backend Folder Structure", height_inches=2.8)

    add_heading(doc, "4.3 MySQL Connection", level=2)
    add_paragraph(
        doc,
        "Database connectivity is implemented in app/core/database.py using SQLAlchemy 2.x. Connection "
        "parameters (DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD) are read from environment "
        "variables. A get_db() dependency provides a scoped session per request with automatic cleanup.",
    )
    add_heading(doc, "4.4 SQLAlchemy Models and Pydantic Schemas", level=2)
    add_paragraph(
        doc,
        "SQLAlchemy models in app/models/entities.py map exactly to existing MySQL tables without "
        "schema modification. Pydantic schemas in app/schemas/entities.py validate all API inputs and "
        "outputs, with separate Create, Update, and Response models for each entity.",
    )
    add_heading(doc, "4.5 CRUD APIs and REST Development", level=2)
    add_paragraph(
        doc,
        "Complete CRUD operations were implemented for all nine tables. Each resource supports POST "
        "(create), GET (list with pagination, search, sorting), GET by ID, PUT (update), and DELETE. "
        "API prefixes include /api/weather, /api/panels, /api/predictions, /api/energy, /api/battery, "
        "/api/battery-status, /api/telemetry, /api/alerts, and /api/logs.",
    )
    add_figure_placeholder(doc, "Swagger UI – API Documentation", height_inches=3.2)
    add_figure_placeholder(doc, "CRUD API Testing via Swagger", height_inches=3.0)
    add_figure_placeholder(doc, "Terminal Output – Backend Server Running", height_inches=2.5)

    add_heading(doc, "4.6 Health and Mock Endpoints", level=2)
    add_paragraph(
        doc,
        "Health endpoints include GET /health (liveness with database status) and GET /health/ready "
        "(readiness probe returning HTTP 503 if MySQL is unavailable). Mock AI endpoints "
        "GET /api/mock/solar-prediction and GET /api/mock/energy return realistic placeholder data "
        "for frontend development until real AI models are integrated by the AI team.",
    )
    add_heading(doc, "4.7 Error Handling and CORS", level=2)
    add_paragraph(
        doc,
        "Centralized exception handlers return structured JSON error responses for validation failures "
        "(422), not-found errors (404), database errors (500), and application errors (400). CORS is "
        "configured for the Next.js frontend running on http://localhost:8501, allowing secure "
        "cross-origin API access during development and deployment.",
    )

    # --- 5. Frontend ---
    add_heading(doc, "5. Frontend Development", level=1)
    add_heading(doc, "5.1 Technology Stack", level=2)
    add_paragraph(
        doc,
        "The dashboard frontend was built using Next.js 15 with the App Router, TypeScript for type "
        "safety, Tailwind CSS for utility-first styling, shadcn/ui for accessible UI components, "
        "Recharts for data visualization, Framer Motion for animations, and Lucide icons for consistent "
        "iconography. The application runs on port 8501.",
    )
    add_heading(doc, "5.2 Dashboard Design and User Experience", level=2)
    add_paragraph(
        doc,
        "I designed an enterprise-grade dark-theme dashboard with a collapsible sidebar, responsive "
        "top navigation bar, and a 12-column grid layout optimized for desktop and mobile viewports. "
        "The interface presents KPI cards, energy flow visualizations, weather intelligence, battery "
        "monitoring, device tables, AI prediction widgets, and system alerts in a cohesive professional "
        "layout suitable for industrial solar monitoring applications.",
    )
    add_paragraph(doc, "Key frontend pages developed include:", bold=False)
    pages = [
        "Main Dashboard (/) – KPIs, solar factory overview, energy flow, and AI widgets",
        "Solar Analytics (/solar-analytics) – generation and performance charts",
        "AI Predictions (/ai-predictions) – solar orientation and consumption forecast views",
        "Battery (/battery) – SOC, health, and charge/discharge history",
        "Energy Flow (/energy-flow) – animated system energy map",
        "Weather (/weather) – weather intelligence and AQI for Delhi",
        "Devices (/devices) – device monitoring table",
        "Reports (/reports) – demo reports with export UI",
        "Settings (/settings) – application preferences stored in localStorage",
    ]
    for page in pages:
        add_bullet(doc, page)

    add_figure_placeholder(doc, "Dashboard – Main Overview Page", height_inches=3.2)
    add_figure_placeholder(doc, "Weather Page", height_inches=3.0)
    add_figure_placeholder(doc, "Battery Page", height_inches=3.0)
    add_figure_placeholder(doc, "Telemetry / Devices Page", height_inches=3.0)
    add_figure_placeholder(doc, "Alerts Page", height_inches=3.0)
    add_figure_placeholder(doc, "System Logs / Reports Page", height_inches=3.0)

    add_heading(doc, "5.3 Frontend–Backend Integration Readiness", level=2)
    add_paragraph(
        doc,
        "The frontend includes api-config.ts with centralized API route definitions pointing to "
        "http://localhost:8000. Dashboard pages currently use structured mock data for demonstration; "
        "the API layer is ready for wiring to live backend endpoints as the integration phase progresses.",
    )

    # --- 6. Challenges ---
    add_heading(doc, "6. Challenges Faced and Solutions", level=1)
    add_heading(doc, "6.1 Database Challenges", level=2)
    add_paragraph(
        doc,
        "Challenge: Aligning SQLAlchemy models with an existing MySQL schema where column names differ "
        "from Python conventions (e.g., MySQL column datetime mapped to recorded_at in the API). "
        "Solution: Used SQLAlchemy mapped_column with explicit column name aliases and validated mappings "
        "against live database columns using inspection scripts and API testing.",
    )
    add_heading(doc, "6.2 Backend Challenges", level=2)
    add_paragraph(
        doc,
        "Challenge: Implementing consistent CRUD logic across nine tables without code duplication. "
        "Solution: Created a generic CRUDBase class and a create_crud_router factory, reducing repetition "
        "while maintaining per-resource customization for alerts and dashboard endpoints.",
    )
    add_paragraph(
        doc,
        "Challenge: IDE import resolution errors for FastAPI and SQLAlchemy packages. "
        "Solution: Configured pyrightconfig.json, VS Code settings, and editable package installation "
        "to point the IDE to backend/.venv with correct PYTHONPATH.",
    )
    add_heading(doc, "6.3 Frontend Challenges", level=2)
    add_paragraph(
        doc,
        "Challenge: Building a responsive enterprise dashboard with consistent dark-theme styling across "
        "nine pages. Solution: Adopted Tailwind CSS design tokens, shadcn/ui component library, and a "
        "shared layout with collapsible sidebar and reusable card/chart components.",
    )
    add_paragraph(
        doc,
        "Challenge: Coordinating frontend development before AI models were ready. "
        "Solution: Implemented mock AI endpoints in the backend and mock data modules in the frontend "
        "so UI development could proceed independently of the AI team.",
    )

    # --- 7. Skills ---
    add_heading(doc, "7. Skills and Technologies Learned", level=1)
    skills = [
        ("MySQL", "Designing relational schemas, EER diagrams, foreign keys, and query validation in MySQL Workbench."),
        ("FastAPI", "Building REST APIs with automatic OpenAPI documentation, dependency injection, and middleware."),
        ("SQLAlchemy", "ORM model mapping, session management, query building, and relationship definitions."),
        ("REST APIs", "Designing resource-oriented endpoints with proper HTTP methods, status codes, and pagination."),
        ("Next.js", "App Router architecture, server/client components, routing, and production builds."),
        ("TypeScript", "Type-safe frontend development with interfaces and strict component typing."),
        ("Tailwind CSS", "Utility-first responsive styling and dark-theme design systems."),
        ("shadcn/ui", "Accessible, composable UI components integrated with Radix UI primitives."),
        ("Git & GitHub", "Branch-based collaboration, commits, pushes, and team workflow on member3-fullstack."),
    ]
    for name, desc in skills:
        add_paragraph(doc, f"{name}: ", bold=True, space_after=0)
        p = doc.paragraphs[-1]
        run = p.add_run(desc)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(6)

    # --- 8. Conclusion ---
    add_heading(doc, "8. Conclusion", level=1)
    add_paragraph(
        doc,
        "Through my role as Member 3 – Full Stack Developer, I successfully delivered the database "
        "connectivity layer, a production-ready FastAPI backend with complete CRUD APIs for all nine "
        "MySQL tables, health and mock endpoints, and a professional Next.js dashboard frontend. My "
        "work provides a solid foundation for AI model integration by Member 1 and Member 2, and cloud "
        "deployment by Member 4.",
    )
    add_paragraph(
        doc,
        "The backend is verified to start successfully, connect to MySQL, expose all routes in Swagger "
        "documentation, and return correct HTTP responses. The frontend presents a modern, responsive "
        "monitoring interface ready for live API integration. This contribution demonstrates practical "
        "full-stack engineering skills including database design, REST API development, frontend UI/UX "
        "design, and collaborative software development using Git and GitHub.",
    )
    add_paragraph(
        doc,
        "Future work includes wiring dashboard pages to live backend APIs and supporting the DevOps team "
        "during production deployment using the handoff documentation provided in backend/HANDOFF.md.",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(f"Report generated: {path}")
